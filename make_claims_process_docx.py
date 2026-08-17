"""
make_claims_process_docx.py - Day 5
Generates a SYNTHETIC "How to File a Claim" Word document.
"""

from docx import Document
from docx.shared import Pt

OUT_PATH = "source_docs/claims_process.docx"


def build_docx():
    doc = Document()

    doc.add_heading("How to File a Claim - SampleCare Health Plans", level=1)
    doc.add_paragraph(
        "This document explains the claims submission and reimbursement process "
        "for SampleCare Health Plans members. This is a synthetic demo document; "
        "it does not describe a real insurance company or real member data."
    )

    doc.add_heading("Step 1: Gather Required Documents", level=2)
    for item in [
        "Itemized bill or receipt from the provider",
        "Proof of payment (if you paid out of pocket)",
        "Your member ID card",
        "Completed claim form (available on the member portal)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Step 2: Submit Your Claim", level=2)
    doc.add_paragraph(
        "Claims can be submitted online through the member portal, by mail, or "
        "by fax. Online submission is fastest and typically processed within "
        "5-7 business days. Mailed claims may take 10-14 business days."
    )

    doc.add_heading("Step 3: Track Your Claim Status", level=2)
    doc.add_paragraph(
        "Once submitted, you can track your claim status in the member portal "
        "under 'My Claims'. Each claim will show one of the following statuses: "
        "Pending, Approved, or Denied."
    )

    doc.add_heading("Step 4: Appeals Process", level=2)
    doc.add_paragraph(
        "If your claim is denied, you have the right to appeal within 180 days "
        "of the denial notice. Appeals can be submitted through the member "
        "portal or by mailing a written appeal to the address listed on your "
        "denial letter. Most appeals are resolved within 30 days."
    )

    doc.add_heading("Frequently Asked Questions", level=2)
    faqs = [
        ("How long does reimbursement take?",
         "Approved claims are typically reimbursed within 14 business days of approval."),
        ("What if I lost my itemized bill?",
         "Contact your provider's billing department to request a duplicate itemized statement."),
        ("Can I submit a claim for a family member?",
         "Yes, as long as they are listed as a dependent on your plan."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    doc.add_paragraph("")
    footer = doc.add_paragraph("Member Services: 1-800-555-0182  |  Synthetic demo document - not real member data")
    footer.runs[0].font.size = Pt(9)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build_docx()